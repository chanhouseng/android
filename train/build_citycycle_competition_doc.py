from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\chanh\Desktop\andorid\website\train\Ai projeect")
OUT = ROOT / "CityCycle_Functionality_Test_Project_Competition_zh-TW.docx"
WIRE = ROOT / "media-files" / "wireframes"

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft JhengHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "25364A"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.append(color)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_real_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n_id = OxmlElement("w:numId")
    n_id.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(n_id)


def add_bullet(doc, text, num_id):
    p = doc.add_paragraph(style="Requirement Bullet")
    apply_numbering(p, num_id)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    return p


def add_wireframe(doc, filename, caption, width_inches=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    shape = run.add_picture(str(WIRE / filename), width=Inches(width_inches))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)


def add_requirements_page(
    doc, num_id, number, title, ui_items, function_items, image_name, image_width=5.0
):
    page_heading = add_heading(doc, f"{number}. \"{title}\" 頁面", 1)
    page_heading.paragraph_format.page_break_before = True
    add_heading(doc, f"{number}.1 介面要求", 2)
    for item in ui_items:
        add_bullet(doc, item, num_id)
    add_heading(doc, f"{number}.2 功能要求", 2)
    for item in function_items:
        add_bullet(doc, item, num_id)
    add_wireframe(
        doc,
        image_name,
        f"WF-{number:02d} — \"{title}\" 頁面線框圖",
        image_width,
    )


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    tr_pr = hdr._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9.5, color=INK)
    set_table_geometry(table, widths_dxa)
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles.add_style("Requirement Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.base_style = normal
    bullet.font.name = FONT_LATIN
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.15


def main():
    doc = Document()
    configure_styles(doc)
    num_id = add_real_bullet_numbering(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("CityCycle Operations | Module A")
    set_font(hr, size=8.5, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    # Cover
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("練習測試題目")
    set_font(r, size=13, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("模組 A：功能實作")
    set_font(r, size=22, bold=True, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('"CityCycle Operations"')
    set_font(r, size=30, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("平板應用程式開發")
    set_font(r, size=15, color=MUTED)
    p.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2.5 小時  |  平板橫向  |  技術中立")
    set_font(r, size=11.5, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(72)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本文件為獨立練習材料，並非 WorldSkills 官方測試題目。")
    set_font(r, size=9.5, italic=True, color=MUTED)

    # Contents
    doc.add_page_break()
    add_heading(doc, "目錄", 1)
    contents = [
        ("簡介", "3"),
        ("專案與任務說明", "3"),
        ("一般要求", "4"),
        ('1. "Dashboard" 頁面', "5"),
        ('2. "Station Management" 頁面', "6"),
        ('3. "Rental Console" 頁面', "7"),
        ('4. "Active Rentals" 頁面', "8"),
        ('5. "Smart Assistant" 頁面', "9"),
        ('6. "Rental History" 頁面', "10"),
        ("參賽者說明", "11"),
    ]
    for label, page in contents:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(6.15))
        r = p.add_run(label + "\t" + page)
        set_font(r, size=11, color=INK)

    # Introduction
    doc.add_page_break()
    add_heading(doc, "簡介", 1)
    p = doc.add_paragraph()
    r = p.add_run(
        "CityCycle 是一套城市共享單車營運系統。營運人員需要掌握站點狀態、管理單車調度、建立與結束租借、監察進行中的租借，以及查閱過往紀錄。"
    )
    set_font(r, size=11, color=INK)
    p = doc.add_paragraph()
    r = p.add_run(
        "你的任務是根據本題要求、所提供的 JSON 資料及線框圖，完成平板版 \"CityCycle Operations\" 應用程式。"
    )
    set_font(r, size=11, color=INK)

    add_heading(doc, "專案與任務說明", 1)
    p = doc.add_paragraph()
    r = p.add_run(
        "應用程式必須在沒有網路連線的情況下完成核心操作。首次啟動時使用所提供的資料建立應用程式內容；使用者在應用程式內產生的變更必須儲存在裝置本機。"
    )
    set_font(r, size=11, color=INK)

    add_heading(doc, "提供檔案", 2)
    add_table(
        doc,
        ["檔案", "內容"],
        [
            ("stations.json", "站點資料"),
            ("bikes.json", "單車資料及目前狀態"),
            ("pricing_rules.json", "租借方案、時間及收費資料"),
            ("members.json", "會員資料"),
            ("active_rentals.json", "進行中的租借資料"),
            ("rental_history.json", "已完成的租借紀錄"),
        ],
        [2700, 6660],
    )

    # General demands
    doc.add_page_break()
    add_heading(doc, "一般要求", 1)
    general = [
        "應用程式必須以平板橫向模式顯示，並在作業系統中使用名稱 \"CityCycle Operations\"。",
        "應用程式必須提供固定導覽，包含 \"Dashboard\"、\"Stations\"、\"Rental Console\"、\"Active Rentals\"、\"Smart Assistant\" 及 \"Rental History\"；目前頁面必須有清楚的選取狀態。",
        "所有頁面及功能必須使用所提供的 JSON 資料，並正確反映應用程式內的最新狀態。",
        "調度、租借、延長、還車及設定的結果，必須在關閉並重新啟動應用程式後保留。",
        "切換頁面、將應用程式移至背景或重新啟動，不得令進行中的租借計時失去正確狀態。",
        "搜尋、篩選、排序、統計及費用顯示必須在相關資料改變後更新。",
        "無法完成的操作或無效輸入必須顯示清楚的應用程式內訊息，且不得改變原有資料。",
        "所有以半形雙引號標示的文字均為程式需要顯示的英文文字，必須依題目指定的拼字及大小寫呈現。",
        "線框圖只用作介面結構與內容配置參考；未在題目中指定的視覺細節可自行設計。",
        "可使用比賽環境允許的任何技術完成題目，但核心功能不得依賴網路服務。",
    ]
    for item in general:
        add_bullet(doc, item, num_id)

    add_requirements_page(
        doc,
        num_id,
        1,
        "Dashboard",
        [
            "頁面標題 \"Dashboard\"。",
            "摘要區域包含 \"Available Bikes\"、\"Empty Docks\"、\"Active Rentals\" 及 \"Overtime\"。",
            "顯示 \"Station Status\" 區域及站點搜尋控制項。",
            "顯示 \"Quick Actions\"，包含 \"Start Rental\"、\"Manage Stations\" 及 \"View Active Rentals\"。",
        ],
        [
            "所有摘要值必須使用目前的站點、單車及租借資料顯示正確結果。",
            "\"Station Status\" 必須顯示站點名稱、區域、可用單車、空車位及目前狀態。",
            "站點搜尋必須即時更新 \"Station Status\" 的顯示結果。",
            "每個 \"Quick Actions\" 操作必須前往相應的應用程式頁面。",
            "調度、開始租借、延長、狀態轉換或還車後，本頁內容必須同步更新。",
        ],
        "WF-01-dashboard.png",
    )

    add_requirements_page(
        doc,
        num_id,
        2,
        "Station Management",
        [
            "頁面標題 \"Station Management\"。",
            "提供搜尋欄，以及 \"District\"、\"Status\"、\"Sort\" 及 \"Reset\" 控制項。",
            "搜尋欄可輸入站點名稱或站點 ID。",
            "\"Stations\" 清單顯示站點 ID、名稱、可用單車、空車位及狀態。",
            "\"Selected Station\" 區域顯示所選站點資料及 \"Bikes assigned to station\" 清單。",
        ],
        [
            "站點的可用單車、空車位及健康狀態必須根據目前資料正確顯示，並隨資料改變更新。",
            "搜尋、篩選及排序必須即時生效；多個篩選條件必須同時套用。",
            "\"Reset\" 必須清除所有搜尋及篩選條件，並恢復預設排序。",
            "選擇不同站點時，\"Selected Station\" 及其單車清單必須更新。",
            "無符合條件的站點時顯示 \"No stations found\"。",
            "使用者可將目前可調度的單車拖放至另一個仍有空車位的站點。",
            "目的站沒有空車位時顯示 \"Destination station is full\"；不可調度的單車顯示 \"Bike is not available for transfer\"。",
            "成功調度後，起點與目的站的內容必須立即更新，並在重新啟動後保留。",
        ],
        "WF-02-station-management.png",
        4.25,
    )

    add_requirements_page(
        doc,
        num_id,
        3,
        "Rental Console",
        [
            "頁面標題 \"Rental Console\"。",
            "\"Rental Setup\" 提供 \"Member\"、\"Bike\"、\"Plan\" 及 \"Add insurance\" 控制項。",
            "\"Rental Summary\" 顯示會員、單車、起點站、解鎖費、基本價格及保險費。",
            "顯示 \"Estimated total\" 及驗證訊息區域。",
            "提供 \"Reset\" 及 \"Start Rental\" 按鈕。",
        ],
        [
            "\"Member\"、\"Bike\" 及 \"Plan\" 必須完成選擇；無法使用的資料必須在驗證訊息區域顯示原因。",
            "\"Bike\" 只可選擇目前可供租借的單車，並在摘要中顯示其起點站。",
            "每位會員最多可同時擁有兩筆未完成租借；超出限制時顯示 \"Rental limit reached\"。",
            "更改會員、單車、方案或保險選擇時，\"Rental Summary\" 及 \"Estimated total\" 必須根據所提供資料更新。",
            "\"Reset\" 必須清除目前輸入、驗證訊息及費用摘要。",
            "開始租借前必須再次確認單車仍可使用；否則顯示 \"Bike is no longer available\"。",
            "成功後顯示 \"Rental started\"，新增租借、更新單車狀態，並切換至 \"Active Rentals\"。",
            "任何驗證失敗均不得建立部分租借或改變單車狀態。",
        ],
        "WF-03-rental-console.png",
        4.25,
    )

    add_requirements_page(
        doc,
        num_id,
        4,
        "Active Rentals",
        [
            "頁面標題及清單標題為 \"Active Rentals\"。",
            "提供搜尋欄、\"Status\"、\"Sort\"、\"Reset\" 及目前租借數量。",
            "每筆租借顯示租借 ID、會員、單車、方案、狀態、剩餘或超時時間、進度及目前費用。",
            "每筆租借提供 \"Extend\" 及 \"Return\" 操作。",
        ],
        [
            "搜尋、狀態篩選及排序必須即時更新清單；\"Reset\" 必須恢復預設顯示。",
            "未超時租借顯示 \"Active\"，剩餘時間及進度必須持續更新。",
            "到達方案時間後必須自動轉為 \"Overtime\"，並持續更新超時時間及費用。",
            "\"Extend\" 必須按所選方案更新租借時間及相關費用；延長後的狀態必須正確反映新的時間。",
            "\"Return\" 必須開啟還車流程，且只可將單車歸還至仍有空車位的站點。",
            "完成還車後，租借狀態、單車位置、站點內容及費用資料必須更新，並將紀錄加入 \"Rental History\"。",
            "應用程式進入背景或重新啟動後，清單必須恢復正確的時間、狀態、進度及費用。",
        ],
        "WF-04-active-rentals.png",
    )

    add_requirements_page(
        doc,
        num_id,
        5,
        "Smart Assistant",
        [
            "頁面標題 \"Smart Assistant\"。",
            "\"Trip Preferences\" 提供 \"Origin\"、\"Destination\" 及推薦方式選擇。",
            "推薦方式包含 \"Nearest bike\"、\"Maximum availability\" 及 \"Lowest estimated cost\"。",
            "提供 \"Find Recommendation\" 按鈕。",
            "\"Recommended Trip\" 顯示起點站、還車站、建議單車、方案、步行距離及預估費用。",
            "顯示警告或無結果訊息區域，以及 \"Apply Recommendation\" 按鈕。",
        ],
        [
            "\"Origin\"、\"Destination\" 及推薦方式必須完成選擇後才可取得推薦。",
            "\"Find Recommendation\" 必須根據目前站點、單車、方案及使用者選擇顯示合適結果。",
            "推薦結果只可使用目前可租借的單車及可完成行程的站點。",
            "選擇不同推薦方式時，結果必須反映該選擇；資料改變後再次搜尋必須使用最新狀態。",
            "無法提供推薦時，必須在訊息區域顯示清楚原因，且不得顯示不可用的結果。",
            "\"Apply Recommendation\" 必須將推薦的單車及方案帶入 \"Rental Console\"，讓使用者繼續建立租借。",
        ],
        "WF-05-smart-assistant.png",
    )

    add_requirements_page(
        doc,
        num_id,
        6,
        "Rental History",
        [
            "頁面標題 \"Rental History\"。",
            "提供租借搜尋欄，以及 \"Date range\"、\"Member\"、\"Status\" 及 \"Reset\" 控制項。",
            "\"Completed Rentals\" 顯示租借紀錄列表及分頁控制。",
            "\"Rental Detail\" 顯示所選租借的完整資料及費用明細。",
        ],
        [
            "列表必須包含所提供的已完成紀錄，以及使用者在應用程式內新完成的租借。",
            "搜尋欄可使用租借、會員或單車資料搜尋；搜尋、日期、會員及狀態條件必須同時套用。",
            "\"Reset\" 必須清除所有搜尋及篩選條件，並返回第一頁。",
            "紀錄必須由最新完成時間開始排列。",
            "列表必須顯示租借、會員、單車、實際時間、狀態及總額，並提供 \"Previous\" 及 \"Next\" 分頁操作。",
            "選擇紀錄後，\"Rental Detail\" 必須顯示租借 ID、起點、終點、實際時間、解鎖費、基本價格、保險費、超時費及總額。",
            "沒有符合條件的紀錄時顯示 \"No rental records\"。",
        ],
        "WF-06-rental-history.png",
    )

    # Submission
    submission_heading = add_heading(doc, "參賽者說明", 1)
    submission_heading.paragraph_format.page_break_before = True
    instructions = [
        "建立套件名稱或組織識別碼 org.citycycle.functionality.xx。",
        "將完整專案儲存在 XX_CityCycle_Functionality 資料夾。",
        "Android 可執行檔命名為 XX_CityCycle_Functionality.apk；iOS 可執行檔命名為 XX_CityCycle_Functionality.app。",
        "將可執行檔放在專案資料夾根目錄，並保留題目提供的六份 JSON 資料。",
        "將完整資料夾提交至指定位置；只有截止時間前完成的版本會被評分。",
        "XX 代表工作站代碼。",
    ]
    for item in instructions:
        add_bullet(doc, item, num_id)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 題目完 —")
    set_font(r, size=11, bold=True, color=BLUE)

    doc.core_properties.title = "CityCycle Operations - Module A Functionality Test Project"
    doc.core_properties.subject = "Independent WorldSkills-style practice test project"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "CityCycle, Functionality, Tablet, Practice"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
